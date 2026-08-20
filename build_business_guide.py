from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\Administrator\Documents\ChatGPT\淘宝文创\docs\淘宝文创业务运营操作指南.docx")
BLUE = "2E74B5"; DARK = "1F4D78"; NAVY = "17365D"; LIGHT = "E8EEF5"; GRAY = "666666"; PALE = "F4F6F9"; RED = "9B1C1C"; GREEN = "276749"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def font(run, size=None, bold=None, color=None, name='微软雅黑'):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)

def para(doc, text='', style=None, bold_prefix=None):
    p=doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); font(r,bold=True); r=p.add_run(text[len(bold_prefix):]); font(r)
    else:
        r=p.add_run(text); font(r)
    return p

def bullet(doc, text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); r=p.add_run(text); font(r); return p

def step(doc, title, detail):
    p=doc.add_paragraph(style='List Number'); r=p.add_run(title+'：'); font(r,bold=True,color=DARK); r=p.add_run(detail); font(r); return p

def callout(doc, title, text, color=BLUE):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.5)
    c=t.cell(0,0); shade(c, PALE); set_cell_margins(c,160,180,160,180)
    p=c.paragraphs[0]; r=p.add_run(title+'  '); font(r,bold=True,color=color); r=p.add_run(text); font(r)
    para(doc,'')

def table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,(h,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Inches(w); shade(c,LIGHT); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; r=p.add_run(h); font(r,bold=True,color=NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i,(v,w) in enumerate(zip(row,widths)):
            cells[i].width=Inches(w); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; r=p.add_run(str(v)); font(r,size=9.5)
    # exact fixed geometry
    tblPr=t._tbl.tblPr; tblW=tblPr.find(qn('w:tblW')); tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    return t

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Title',28,NAVY,0,8),('Subtitle',13,GRAY,0,18),('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK,10,5)]:
    s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=(name!='Subtitle')
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for lname in ['List Bullet','List Bullet 2','List Number']:
    s=styles[lname]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); s.font.size=Pt(11); s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.25

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=hp.add_run('淘宝文创 · 业务运营手册'); font(r,size=9,color=GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=fp.add_run('内部操作参考｜版本 1.0｜2026-08-20'); font(r,size=8.5,color=GRAY)

# Cover
para(doc,'业务运营操作指南',style='Title').alignment=WD_ALIGN_PARAGRAPH.CENTER
para(doc,'淘宝文创商品采集、筛选、图片处理与淘宝刊登',style='Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
para(doc,'').paragraph_format.space_after=Pt(70)
p=para(doc,'适用对象'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.runs[0],size=10,bold=True,color=BLUE)
p=para(doc,'选品运营 · 商品运营 · 视觉运营 · 店铺运营 · 业务负责人'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.runs[0],size=12,color=NAVY)
para(doc,'').paragraph_format.space_after=Pt(40)
callout(doc,'使用目标','让业务人员按统一流程完成“采集 → 入库 → 筛选 → 图片处理 → 人工审核 → 淘宝刊登 → 结果回查”，并在异常时知道应检查什么、交给谁处理。')
table(doc,['文档信息','内容'],[['系统入口','http://127.0.0.1:8000'],['业务范围','1688/淘宝商品采集；本地草稿审核；图片任务；淘宝/天猫刊登'],['操作原则','先审核、后处理；先预览、后刊登；所有关键步骤保留任务记录'],['当前边界','1688 用作供货来源；淘宝/天猫用作销售刊登目标']], [1.6,4.9])

doc.add_page_break()
para(doc,'目录与快速入口',style='Heading 1')
for x in ['1. 使用前准备与角色分工','2. 系统登录与首页检查','3. 店铺与平台配置','4. 商品采集','5. 商品草稿筛选与入库','6. 图片处理与人工审核','7. 淘宝刊登','8. 任务查询、恢复与异常处理','9. 日常运营检查清单','10. 常见问题与业务口径']:
    bullet(doc,x)
callout(doc,'一条主流程','采集商品 → 在草稿库核对 → 剔除/保留 → 创建图片任务 → 审核成图 → 创建刊登任务 → 人工确认 → 回查商品 ID 与链接。',GREEN)

para(doc,'1. 使用前准备与角色分工',style='Heading 1')
table(doc,['角色','主要工作','不应跳过'],[
 ['选品运营','提交采集、核对来源商品、剔除重复或不合格商品','来源链接、价格、SKU、图片完整性'],
 ['商品运营','完善标题、卖点、类目、属性、价格和库存','类目与 Schema 校验'],
 ['视觉运营','发起图片处理、选择模板、检查主图与详情图','尺寸、文字、版权与输出地址'],
 ['店铺运营','选择店铺、创建刊登任务、回查上架结果','店铺授权、最终预览、平台结果'],
 ['管理员','维护平台密钥、店铺授权、服务状态','敏感凭据不得发群或写入文档']], [1.1,2.65,2.75])
para(doc,'准备信息',style='Heading 2')
for x in ['管理端地址及个人账号。','淘宝目标店铺已在系统中建立并完成授权测试。','目标类目 ID；天猫商品还需对应 SPU ID。','类目对应的淘宝发布 Schema XML 模板。','图片模板、输出尺寸、文字规范和审核标准。']:
    bullet(doc,x)

para(doc,'2. 系统登录与首页检查',style='Heading 1')
step(doc,'打开管理端','浏览器访问 http://127.0.0.1:8000，输入管理员分配的账号。')
step(doc,'确认服务可用','页面菜单可以正常打开；提交操作时没有 502/503 提示。')
step(doc,'确认平台入口','在“店铺管理 → 新建店铺”中应看到“淘宝 / 天猫开放平台”和“1688 供货店铺”。')
callout(doc,'登录信息','账号、密码、App Key、App Secret、Session/Access Token 由管理员单独分配。首次登录后应按团队制度修改密码；凭据不要粘贴到商品备注或业务群。',RED)

para(doc,'3. 店铺与平台配置',style='Heading 1')
para(doc,'3.1 淘宝/天猫销售店铺',style='Heading 2')
step(doc,'配置平台','设置 → 平台接入设置 → 淘宝开放平台，填写 App Key、App Secret；沙箱和自定义网关仅在管理员确认后使用。')
step(doc,'新建店铺','店铺管理 → 新建店铺，销售平台选择“淘宝 / 天猫开放平台（taobao）”。')
step(doc,'授权店铺','进入店铺授权，填写 Session/Access Token。')
step(doc,'测试连接','点击“测试连接”；成功后再安排刊登。')
para(doc,'3.2 1688 供货店铺',style='Heading 2')
step(doc,'登记来源','店铺管理 → 新建店铺，选择“1688 供货店铺（1688）”。')
step(doc,'填写资料','填写店铺名称，可选填供应商或会员 ID。')
callout(doc,'平台定位','1688 当前作为供货来源与采集入口，不作为销售刊登目标；成品发布到已授权的淘宝/天猫店铺。')

para(doc,'4. 商品采集',style='Heading 1')
para(doc,'4.1 按商品链接采集',style='Heading 2')
step(doc,'进入任务页','采集 → 采集任务。')
step(doc,'选择平台','淘宝/天猫商品选“淘宝/天猫采集器”；1688 供货商品选 1688。')
step(doc,'粘贴链接','使用标准商品详情页链接，避免短链、活动跳转页或已失效链接。')
step(doc,'提交任务','点击提交后，在任务列表查看进度和结果。')
step(doc,'处理登录验证','出现登录或安全验证提示时，打开对应采集浏览器，人工完成验证后重新提交。')
para(doc,'4.2 采集完成检查',style='Heading 2')
table(doc,['检查项','合格标准'],[['标题','商品名称清晰，无乱码或无关促销词'],['价格','与来源页面一致，币种和小数位正常'],['图片','主图、详情图均可打开，无明显水印或失效链接'],['规格/SKU','颜色、尺寸、价格、库存关系完整'],['参数','材质、尺寸、品牌/产地等关键字段齐全'],['来源','来源平台与原始链接可追溯']], [1.35,5.15])

para(doc,'5. 商品草稿筛选与入库',style='Heading 1')
step(doc,'打开商品草稿库','按采集时间、平台、状态筛选本批商品。')
step(doc,'先剔除','删除或标记重复商品、信息缺失、图片质量差、价格无优势、版权风险高的商品。')
step(doc,'再保留','对候选商品补齐业务标签、目标店铺、目标类目和负责人。')
step(doc,'完善销售信息','调整标题、卖点、描述、售价、库存与 SKU 编码。')
step(doc,'保存草稿','确认来源数据仍可追溯，再进入图片处理。')
callout(doc,'选品建议','不要直接将采集结果批量刊登。建议使用“来源可追溯、毛利达标、图片可加工、类目可发布、售后风险可控”五项门槛。',GREEN)

para(doc,'6. 图片处理与人工审核',style='Heading 1')
step(doc,'选择商品','只对已标记“保留”的商品创建图片任务。')
step(doc,'选择处理方式','选择系统中已配置的图片 Provider、模板或 Photoshop API 流程。')
step(doc,'填写规格','明确主图/详情图、画布尺寸、背景、品牌元素、文字内容和输出格式。')
step(doc,'提交并等待','图片处理是异步任务；不要连续重复点击提交。')
step(doc,'人工审核','逐张检查构图、清晰度、文字、尺寸、商品一致性、平台规则和输出链接。')
step(doc,'确认成图','合格后标记为待刊登；不合格时记录原因并重新处理。')
table(doc,['审核项','运营判断'],[['商品一致性','颜色、款式、数量、配件与真实商品一致'],['文字','无错别字、夸大表述、遮挡主体或超出安全区'],['尺寸','符合店铺当前主图和详情页规范'],['画面','清晰、无拉伸、无异常边缘、无无关水印'],['可用性','图片地址稳定，刊登预览中可正常加载']], [1.45,5.05])
callout(doc,'当前说明','系统已经具备通用图片任务与人工审核链路；具体 Photoshop 自动化能力取决于管理员实际配置的 Provider、模板和接口状态。')

para(doc,'7. 淘宝刊登',style='Heading 1')
para(doc,'7.1 刊登前配置',style='Heading 2')
table(doc,['参数','业务说明'],[['发布市场','taobao、tmall 或 litetao'],['默认类目 ID','目标淘宝类目；任务中可覆盖'],['默认 SPU ID','天猫通常必填；淘宝可留空'],['Biz Type','仅在平台或管理员明确要求时填写'],['Schema XML 模板','对应目标类目的完整发布 Schema']], [1.55,4.95])
para(doc,'可用模板变量：{{title}}、{{description}}、{{main_image_url}}、{{price}}、{{stock}}、{{sku_code}}。系统提交前会执行 XML 转义。')
para(doc,'7.2 创建刊登任务',style='Heading 2')
step(doc,'完成刊登检查','确认商品状态、图片、标题、描述、类目、SKU、售价和库存。')
step(doc,'选择目标店铺','只能选择已授权且测试连接成功的淘宝/天猫店铺。')
step(doc,'预览发布信息','重点核对类目、市场、图片、价格、库存和 Schema。')
step(doc,'创建任务','提交刊登任务并完成人工确认。')
step(doc,'查看结果','成功后记录淘宝商品 ID、商品链接和 publication 状态。')
callout(doc,'重要','刊登接口成功不等于页面展示完全符合预期。运营人员还应打开返回的商品链接，检查前台标题、图片、规格、价格和库存。远端 SKU ID 需后续详情同步绑定。',RED)

para(doc,'8. 任务查询、恢复与异常处理',style='Heading 1')
para(doc,'系统会保存稳定任务 ID、步骤结果和失败记录。重试时从最后已确认状态继续，图片处理与刊登任务分别保存。',style=None)
table(doc,['现象','业务处理','升级给'],[
 ['502/503','记录时间和操作页面；稍后刷新；确认不是重复提交','管理员检查管理端、后端、采集器、数据库和 Redis'],
 ['采集要求登录','打开对应采集浏览器，完成登录/验证后重试','需要账号时联系管理员'],
 ['店铺下拉无淘宝/1688','Ctrl+F5 强制刷新；重新登录','管理员确认后端已更新并重启'],
 ['淘宝连接失败','核对店铺与授权状态，不在群里发送 Token','管理员检查应用权限和凭据'],
 ['缺少类目/SPU','暂停刊登，补充类目；天猫补充 SPU','商品负责人/管理员'],
 ['Schema XML 错误','核对该类目完整 Schema 和字段规则','商品负责人/技术'],
 ['图片处理失败','查看任务原因，修正模板、素材或参数后重试','视觉负责人/管理员'],
 ['刊登成功但前台异常','保留商品 ID 和截图，暂停批量任务','店铺负责人']], [1.4,3.25,1.85])
callout(doc,'避免重复','按钮点击后先看任务列表。只在任务明确失败且原因已处理后重试；不要用新任务绕过尚未确认的旧任务。')

para(doc,'9. 日常运营检查清单',style='Heading 1')
para(doc,'班前',style='Heading 2')
for x in ['□ 管理端可登录，菜单和列表可打开。','□ 淘宝目标店铺授权状态正常。','□ 今日使用的类目、Schema 与图片模板已确认。','□ 待处理失败任务已交接。']:
    bullet(doc,x)
para(doc,'刊登前',style='Heading 2')
for x in ['□ 商品已人工筛选并标记保留。','□ 标题、描述、类目、价格、库存、SKU 均已复核。','□ 主图与详情图已人工审核。','□ 目标店铺、发布市场、类目 ID/SPU ID 正确。','□ 已查看最终预览并完成确认。']:
    bullet(doc,x)
para(doc,'班后',style='Heading 2')
for x in ['□ 当日刊登成功数、失败数已登记。','□ 成功商品已打开前台链接复查。','□ 失败任务已记录任务 ID、错误信息和负责人。','□ 未完成任务已交接，不重复创建。']:
    bullet(doc,x)

para(doc,'10. 常见问题与业务口径',style='Heading 1')
table(doc,['问题','统一回答'],[
 ['支持淘宝吗？','支持淘宝/天猫销售店铺接入与商品刊登，需有效开放平台配置、店铺授权及类目 Schema。'],
 ['支持 1688 上架吗？','当前 1688 用作供货来源登记和商品采集，不作为销售刊登目标。'],
 ['支持图片搜索或关键词搜索吗？','以管理端实际启用的采集器能力为准；链接采集是当前明确可用流程。'],
 ['图片是否完全自动 PS？','系统有图片任务与审核流程；自动化程度取决于已配置的图片 Provider、模板和 Photoshop 接口。'],
 ['刊登后为什么没有远端 SKU ID？','当前发布接口不返回远端 SKU ID，系统保留商品 ID，后续通过详情同步绑定。'],
 ['失败后能否重试？','先处理错误原因，再从任务记录重试；系统会利用稳定任务 ID 和已完成步骤减少重复执行。']], [2.05,4.45])

para(doc,'附录：提交异常工单时的信息',style='Heading 1')
for x in ['发生时间（精确到分钟）','操作账号与页面名称','任务 ID / 商品草稿 ID / 店铺名称','操作步骤与期望结果','页面完整错误提示','必要截图（注意遮挡账号、Token、密钥）','是否重试及重试结果']:
    bullet(doc,x)
callout(doc,'交付标准','业务问题以任务 ID 和商品 ID 为主线；敏感凭据只通过管理员规定的安全渠道处理。')

# Keep table headers together/repeat and prevent split rows where possible
for t in doc.tables:
    trPr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); trPr.append(rep)
    for row in t.rows:
        pr=row._tr.get_or_add_trPr(); cant=OxmlElement('w:cantSplit'); pr.append(cant)

doc.core_properties.title='淘宝文创业务运营操作指南'
doc.core_properties.subject='业务与运营人员操作手册'
doc.core_properties.author='淘宝文创项目组'
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT)
print(OUT)
